import os
from langchain_community.llms import OpenAI
from dotenv import load_dotenv
import logging
import re
import json
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import time
import tempfile
import subprocess
from bs4 import BeautifulSoup

logging.basicConfig(level=logging.INFO)
load_dotenv()

class PromptCreator:
    def __init__(self):
        self.openai = self.init_models()

    def init_models(self):
        api_key = os.getenv("OPENAI_API_KEY")
        model_name = os.getenv('OPENAI_MODEL_NAME', 'gpt-4o-mini')
        logging.info(f"Using model: {model_name}")
        logging.info(f"Using api_key: {api_key}")
        openai = OpenAI(temperature=0.4, api_key=api_key, model_name=model_name)
        return openai

    def run_openai_prompt(self, prompt):
        return self.openai.invoke(prompt)
    
class JSONParserForLLM:

    @staticmethod
    def extract_field(text, field_name):
        """Extract the value of a specific field from the text."""
        # Define a more specific pattern to match the field, handling multi-line values and lists
        pattern = rf'"{field_name}":\s*(\[.*?\]|\{{.*?\}}|".*?"|\d+|true|false|null)(,|\n|$)'
        match = re.search(pattern, text, re.DOTALL)

        if not match:
            return None

        extracted_value = match.group(1).strip()

        # Try to parse the extracted value as JSON
        try:
            parsed_value = json.loads(extracted_value)
        except json.JSONDecodeError:
            # If it's not valid JSON, return the string as is, but strip quotes
            parsed_value = extracted_value.strip('"')

        return parsed_value
    
    @staticmethod
    def clean_json_text(text):
        cleaned_text = re.sub(r"```json|```", "", text).strip()
        output = json.loads(cleaned_text)
        return output
    
    ##





from docx import Document
from docx.enum.section import WD_SECTION
from ai.llm import PromptCreator
import time
import os

def translate_doc_preserving_format(input_doc_path, output_doc_path):
    # Initialize the translator
    prompt_creator = PromptCreator()
    
    # Load the document
    original_doc = Document(input_doc_path)
    
    # Create output document
    output_doc = Document()
    
    # Copy document properties and styles from original
    # This preserves section settings, margins, etc.
    for section in original_doc.sections:
        new_section = output_doc.add_section()
        new_section.start_type = section.start_type
        new_section.page_height = section.page_height
        new_section.page_width = section.page_width
        new_section.left_margin = section.left_margin
        new_section.right_margin = section.right_margin
        new_section.top_margin = section.top_margin
        new_section.bottom_margin = section.bottom_margin
    
    # Collect paragraphs by page
    pages = []
    current_page = []
    
    for para in original_doc.paragraphs:
        # Store the paragraph text and its runs (formatted pieces)
        para_info = {
            'text': para.text,
            'style': para.style.name,
            'alignment': para.alignment,
            'runs': []
        }
        
        # Store information about each run (formatting chunk)
        for run in para.runs:
            run_info = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_size': run.font.size if run.font.size else None,
                'font_name': run.font.name if run.font.name else None,
                'has_break': any(br.type == 'page' for br in run._element.br_lst)
            }
            para_info['runs'].append(run_info)
        
        current_page.append(para_info)
        
        # Check if paragraph or any of its runs has a page break
        if any(run['has_break'] for run in para_info['runs']):
            pages.append(current_page)
            current_page = []
    
    # Add the last page
    if current_page:
        pages.append(current_page)
    
    # If no page breaks were detected, estimate based on paragraphs
    if len(pages) < 2:
        # Fallback to the original method
        pages = []
        paragraphs = [{'text': p.text, 'style': p.style.name, 'alignment': p.alignment} 
                     for p in original_doc.paragraphs if p.text.strip()]
        
        # Estimate paragraphs per page
        page_size = max(1, len(paragraphs) // 80)
        
        for i in range(0, len(paragraphs), page_size):
            pages.append(paragraphs[i:i+page_size])
    
    print(f"Document split into {len(pages)} pages with formatting preserved")
    
    # Translate each page
    for i, page_paras in enumerate(pages[:3]):
        print(f"Translating page {i+1}/{len(pages)}...")
        
        # Extract just the text for translation
        page_text = "\n".join([para['text'] for para in page_paras if para['text'].strip()])
        
        # Skip empty pages
        if not page_text.strip():
            continue
        
        prompt = f"""
        Translate the following text from Hebrew to Russian with high quality and accuracy.
        Preserve ALL paragraph breaks exactly as in the original text.
        Ensure the translation is natural and fluent in Russian.
        Return ONLY the translated text with no additional explanations.
        
        TEXT TO TRANSLATE:
        {page_text}
        """
        
        try:
            # Get translated text
            translation = prompt_creator.run_openai_prompt(prompt)
            # Clean up any markdown
            translation = translation.replace("```", "").strip()
            
            # Split translation back into paragraphs
            translated_paragraphs = translation.split('\n')
            
            # Ensure we have enough translated paragraphs
            # If not, pad with empty strings
            while len(translated_paragraphs) < len([p for p in page_paras if p['text'].strip()]):
                translated_paragraphs.append("")
            
            # Add translated paragraphs to document with original formatting
            para_index = 0
            for para_info in page_paras:
                # Skip empty paragraphs but preserve them in the document
                if not para_info['text'].strip():
                    p = output_doc.add_paragraph()
                    p.style = para_info['style']
                    p.alignment = para_info['alignment']
                    continue
                
                # Get the translated text for this paragraph
                if para_index < len(translated_paragraphs):
                    translated_text = translated_paragraphs[para_index]
                    para_index += 1
                else:
                    translated_text = ""
                
                # Create paragraph with original style
                p = output_doc.add_paragraph()
                p.style = para_info['style']
                p.alignment = para_info['alignment']
                
                # If we have detailed run info, try to map translation to runs
                if 'runs' in para_info and para_info['runs']:
                    # Simple approach: add one run with translation
                    run = p.add_run(translated_text)
                    
                    # Apply formatting from the first run
                    first_run = para_info['runs'][0]
                    run.bold = first_run['bold']
                    run.italic = first_run['italic']
                    run.underline = first_run['underline']
                    if first_run['font_size']:
                        run.font.size = first_run['font_size']
                    if first_run['font_name']:
                        run.font.name = first_run['font_name']
                else:
                    # Simple paragraph without formatting info
                    p.add_run(translated_text)
            
            # Add page break after each page except the last
            if i < len(pages) - 1:
                output_doc.add_page_break()
                
            print(f"✓ Page {i+1} translated successfully with formatting preserved")
            
        except Exception as e:
            print(f"✗ Error translating page {i+1}: {str(e)}")
            # Keep original text with formatting
            for para_info in page_paras:
                p = output_doc.add_paragraph()
                p.style = para_info['style']
                p.alignment = para_info['alignment']
                p.add_run(f"[TRANSLATION ERROR] {para_info['text']}")
            
            # Add page break after each page except the last
            if i < len(pages) - 1:
                output_doc.add_page_break()
        
        # Add delay to avoid rate limits
        time.sleep(1.5)
    
    # Save the translated document
    output_doc.save(output_doc_path)
    print(f"\n✓ Translation complete with formatting preserved! Document saved to: {output_doc_path}")
    return output_doc_path

def translate_doc_exact_format(input_doc_path, output_doc_path):
    """Preserves exact formatting by working with the underlying XML structure"""
    prompt_creator = PromptCreator()
    
    # Step 1: Create a copy of the original document to preserve ALL formatting
    from shutil import copyfile
    copyfile(input_doc_path, output_doc_path)
    
    # Step 2: Open the copy for modification
    doc = docx.Document(output_doc_path)
    
    # Step 3: Extract text while preserving format markers
    text_with_markers = []
    marker_map = {}
    marker_id = 0
    
    # Process paragraphs and preserve special formatting with markers
    for para in doc.paragraphs:
        para_text = ""
        for run in para.runs:
            if run.text.strip():
                # Create a unique marker for this text run
                marker = f"[[MARKER_{marker_id}]]"
                marker_map[marker] = {
                    'paragraph_index': doc.paragraphs.index(para),
                    'run_index': para.runs.index(run),
                    'original_text': run.text
                }
                para_text += marker
                marker_id += 1
        
        if para_text:
            text_with_markers.append(para_text)
    
    # Step 4: Combine text for translation while preserving markers
    text_to_translate = "\n".join(text_with_markers)
    
    # Step 5: Create a clean version for translation (replace markers with actual text)
    clean_text = text_to_translate
    for marker, info in marker_map.items():
        clean_text = clean_text.replace(marker, info['original_text'])
    
    # Step 6: Translate the clean text
    prompt = f"""
    Translate the following Hebrew text to Russian with high quality and accuracy.
    Preserve ALL paragraph breaks exactly as in the original text.
    Ensure the translation is natural and fluent in Russian.
    Return ONLY the translated text with no additional explanations.
    
    TEXT TO TRANSLATE:
    {clean_text}
    """
    
    translated_text = prompt_creator.run_openai_prompt(prompt)
    
    # Step 7: Split the translation by paragraphs
    translated_paragraphs = translated_text.split('\n')
    
    # Step 8: Align translations with original markers
    translation_map = {}
    
    # First create a mapping from original text to translation
    original_paragraphs = clean_text.split('\n')
    min_len = min(len(original_paragraphs), len(translated_paragraphs))
    
    for i in range(min_len):
        # Use regex to identify words in both languages
        original_words = re.findall(r'\S+', original_paragraphs[i])
        translated_words = re.findall(r'\S+', translated_paragraphs[i])
        
        # Map approximate proportions - this is an estimation
        if len(original_words) > 0 and len(translated_words) > 0:
            ratio = len(translated_words) / len(original_words)
            
            for marker in marker_map:
                if marker in text_with_markers[i] if i < len(text_with_markers) else False:
                    original_text = marker_map[marker]['original_text']
                    # Estimate translated length
                    est_length = max(1, int(len(original_text.split()) * ratio))
                    
                    # Get a proportional segment of the translation
                    word_index = original_paragraphs[i].index(original_text) / len(original_paragraphs[i])
                    start_pos = int(word_index * len(translated_words))
                    segment = " ".join(translated_words[start_pos:start_pos + est_length])
                    
                    translation_map[marker] = segment
    
    # Step 9: Replace text in the original document with translations
    for marker, info in marker_map.items():
        para = doc.paragraphs[info['paragraph_index']]
        run = para.runs[info['run_index']]
        
        # Replace with translation if available, otherwise keep original
        if marker in translation_map:
            run.text = translation_map[marker]
        else:
            # If mapping failed, use a simpler approach
            run.text = f"[Translation failed for: {run.text}]"
    
    # Step 10: Save the modified document
    doc.save(output_doc_path)
    return output_doc_path

def translate_with_pandoc(input_doc_path, output_doc_path):
    """Uses pandoc to convert to HTML, translate the text only, then convert back"""
    prompt_creator = PromptCreator()
    
    # Step 1: Convert DOCX to HTML to preserve all formatting
    html_file = tempfile.NamedTemporaryFile(suffix='.html', delete=False).name
    subprocess.run(['pandoc', input_doc_path, '-o', html_file])
    
    # Step 2: Parse HTML to separate text from formatting
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Step 3: Extract text elements while preserving their positions
    text_elements = []
    for element in soup.find_all(text=True):
        if element.strip():
            text_elements.append({
                'element': element,
                'text': element.string
            })
    
    # Step 4: Translate text elements while preserving HTML structure
    for item in text_elements:
        if item['text'].strip():
            prompt = f"""
            Translate only this Hebrew text to Russian:
            {item['text']}
            """
            translation = prompt_creator.run_openai_prompt(prompt).strip()
            # Replace the original text with translation
            item['element'].string.replace_with(translation)
    
    # Step 5: Write updated HTML
    with open(html_file, 'w', encoding='utf-8') as f:
        f.write(str(soup))
    
    # Step 6: Convert HTML back to DOCX
    subprocess.run(['pandoc', html_file, '-o', output_doc_path])
    
    # Clean up temp file
    os.unlink(html_file)
    
    return output_doc_path

# Example usage
if __name__ == "__main__":
    input_file = "/home/ubuntu/pr_repo/survey/SurveyAnalyzerBe/ai/heb.docx"
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' does not exist.")
        exit()
    output_file = "russian_document.docx"
    
    translate_doc_preserving_format(input_file, output_file)